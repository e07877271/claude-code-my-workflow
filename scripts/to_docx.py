#!/usr/bin/env python3
"""
将 reports/ 下的 Markdown 报告转换为 Word (.docx)。
格式对标国内经济金融研究机构报告：仿宋四号正文，宋体三号标题。

用法：
  python scripts/to_docx.py reports/2026-05-02_美元国际地位.md
  python scripts/to_docx.py reports/xxx.md C:/Users/xxx/Desktop/报告.docx
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("缺少依赖，请运行：pip install python-docx")
    sys.exit(1)


# ── 字体辅助 ────────────────────────────────────────────────────
def set_font(run, zh="仿宋", en="Times New Roman", size=14, bold=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), zh)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)


def set_para_indent(para, first_line_chars=2, line_spacing=28):
    """首行缩进 2 字符，固定行距 28 磅"""
    pf = para.paragraph_format
    pf.first_line_indent = Pt(line_spacing * first_line_chars / 2)
    from docx.oxml import OxmlElement
    pPr = para._element.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(int(line_spacing * 20)))  # 磅 → twip
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)


# ── 行内加粗解析 ────────────────────────────────────────────────
def add_inline(para, text, zh="仿宋", en="Times New Roman", size=14):
    """解析 **加粗** 标记，逐段添加 run"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            set_font(r, zh, en, size, bold=True)
        elif part:
            r = para.add_run(part)
            set_font(r, zh, en, size, bold=False)


# ── 主转换函数 ──────────────────────────────────────────────────
def convert(md_path: str, out_path: str):
    doc = Document()

    # 页面：A4，四边 2.5 cm
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2.5))

    # 清除默认段落
    for elem in list(doc.element.body):
        doc.element.body.remove(elem)

    lines = Path(md_path).read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1

        # 空行 / 分隔线 → 跳过
        if not line or line == "---":
            continue

        # ── 总标题（# 开头）──
        if line.startswith("# "):
            title = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(title)
            set_font(r, "宋体", "Times New Roman", 16, bold=True)
            continue

        # ── 一级节标题（汉字数字＋顿号，或 ## 开头）──
        if re.match(r"^[一二三四五六七八九十]+、", line) or line.startswith("## "):
            title = re.sub(r"^##\s*", "", line)
            p = doc.add_paragraph()
            r = p.add_run(title)
            set_font(r, "宋体", "Times New Roman", 14, bold=True)
            pf = p.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after  = Pt(4)
            continue

        # ── 编者按 ──
        if re.match(r"^(编者按[：:]?|\*\*编者按\*\*)", line):
            p = doc.add_paragraph()
            set_para_indent(p)
            text = re.sub(r"^\*?\*?编者按\*?\*?[：:]?\s*", "编者按　", line)
            add_inline(p, text)
            continue

        # ── 表格（简单跳过，不渲染）──
        if line.startswith("|"):
            continue

        # ── 代码块 ──
        if line.startswith("```"):
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1  # 跳过结束 ```
            continue

        # ── 普通正文 ──
        p = doc.add_paragraph()
        set_para_indent(p)
        add_inline(p, line)

    doc.save(out_path)
    print(f"✓ 已保存：{out_path}")


# ── 入口 ────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法：python scripts/to_docx.py <报告.md> [输出路径.docx]")
        sys.exit(1)

    md  = sys.argv[1]
    if len(sys.argv) >= 3:
        out = sys.argv[2]
    else:
        desktop = Path.home() / "Desktop"
        stem = Path(md).stem
        out  = str(desktop / f"{stem}.docx")

    convert(md, out)
